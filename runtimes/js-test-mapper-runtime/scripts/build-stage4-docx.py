#!/usr/bin/env python3
import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"

FONT_CANDIDATES = [
    ("Microsoft YaHei", [Path("C:/Windows/Fonts/msyh.ttc"), Path("C:/Windows/Fonts/msyhbd.ttc")]),
    ("SimHei", [Path("C:/Windows/Fonts/simhei.ttf")]),
    ("SimSun", [Path("C:/Windows/Fonts/simsun.ttc")]),
    ("DengXian", [Path("C:/Windows/Fonts/Deng.ttf")]),
    ("Noto Sans CJK SC", [Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc")]),
]


def choose_cjk_font():
    for family, paths in FONT_CANDIDATES:
        if any(path.exists() for path in paths):
            return family
    return "Arial"


CJK_FONT = choose_cjk_font()


def set_run_font(run, size=11, color=INK, bold=None, italic=None, name=None):
    name = name or CJK_FONT
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), CJK_FONT)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        item = margins.find(qn(f"w:{name}"))
        if item is None:
            item = OxmlElement(f"w:{name}")
            margins.append(item)
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    old_grid = table._tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        table._tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    table._tbl.insert(1, grid)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell_width(cell, widths[index])
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = borders.find(qn(f"w:{name}"))
        if item is None:
            item = OxmlElement(f"w:{name}")
            borders.append(item)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "4")
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), "D9E2F3")


def paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = CJK_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), CJK_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), CJK_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (("Heading 1", 16, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6), ("Heading 3", 12, DARK_BLUE, 8, 4)):
        style = doc.styles[name]
        style.font.name = CJK_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = CJK_FONT
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:ascii"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167


def add_header_footer(doc):
    header = doc.sections[0].header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("无需求-Web JS逆向测试建图")
    set_run_font(run, size=9, color=MUTED, bold=True)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("静态恢复 / 运行观察 / 待执行验证 | ")
    set_run_font(run, size=9, color=MUTED)
    for field_type, text in (("begin", None), ("instr", "PAGE"), ("end", None)):
        element = OxmlElement("w:fldChar" if field_type != "instr" else "w:instrText")
        if field_type == "instr":
            element.set(qn("xml:space"), "preserve")
            element.text = text
        else:
            element.set(qn("w:fldCharType"), field_type)
        run._r.append(element)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(str(text))
    set_run_font(run)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table_geometry(table, widths)
    table_borders(table)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, DARK_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(header))
        set_run_font(run, size=9.5, color=WHITE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_run_font(run, size=9.5)
    table_geometry(table, widths)


def build(input_path, output_path):
    projection = json.loads(Path(input_path).read_text(encoding="utf-8"))
    doc = Document()
    configure(doc)
    add_header_footer(doc)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run(projection.get("title") or "过程小结"), size=24, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(subtitle.add_run(projection.get("subtitle") or "基于 run-data 的测试认知摘要"), size=13, color=MUTED)
    for label, value in (("批次", projection.get("current_batch") or "当前批次"), ("状态", projection.get("status_label") or "静态恢复 / 待执行验证"), ("run-data fingerprint", projection["run_data_fingerprint"])):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(f"{label}: "), size=10.5, color=INK, bold=True)
        set_run_font(p.add_run(str(value)), size=10.5)
    callout = doc.add_paragraph()
    callout.paragraph_format.space_before = Pt(12)
    callout.paragraph_format.space_after = Pt(12)
    callout.paragraph_format.left_indent = Inches(0.12)
    callout.paragraph_format.right_indent = Inches(0.12)
    paragraph_shading(callout, CALLOUT)
    set_run_font(callout.add_run(projection.get("summary", {}).get("lead") or projection.get("summary", {}).get("system") or "以下内容只展示已有证据支持的代表性结构；没有证据的业务含义保持待确认。"), bold=True)

    for chapter in projection["chapters"]:
        doc.add_heading(f"{chapter['number']} {chapter['title']}", level=1)
        for paragraph in chapter.get("paragraphs", []):
            p = doc.add_paragraph()
            set_run_font(p.add_run(paragraph))
        if chapter["number"] not in (4, 5) or not (chapter.get("rules") or chapter.get("risks")):
            for bullet in chapter.get("bullets", []):
                add_bullet(doc, bullet)
        if chapter["number"] == 2:
            rows = []
            for item in chapter.get("bullets", []):
                left, separator, right = item.partition("：")
                rows.append([left if separator else item, right if separator else ""]) 
            add_table(doc, ["功能 / 路由", "当前功能地图"], rows, [3000, 6360])
        elif chapter["number"] == 3:
            rows = [[chain.get("display_id", "待确认"), chain.get("action", chain.get("action_label", "待确认")), chain.get("summary", "证据不足"), chain.get("current_status", "待执行验证")] for chain in chapter.get("chains", [])]
            add_table(doc, ["调用链ID", "业务动作", "代表性链路", "当前状态"], rows, [1300, 1800, 5260, 1000])
        elif chapter["number"] == 4:
            rows = [[item.get("display_id", "待确认"), item.get("type", "规则线索"), item.get("content", ""), item.get("current_status", "待执行验证")] for item in chapter.get("rules", [])]
            add_table(doc, ["规则ID", "类型", "关键规则 / 对测试影响", "当前状态"], rows, [1300, 1200, 5860, 1000])
        elif chapter["number"] == 5:
            rows = [[item.get("display_id", "待确认"), item.get("type", "待确认"), item.get("content", ""), item.get("current_status", "待执行验证")] for item in chapter.get("risks", [])]
            add_table(doc, ["编号", "类型", "关注内容", "当前状态"], rows, [1300, 1200, 5860, 1000])

    props = doc.core_properties
    props.title = "过程小结"
    props.subject = "无需求-Web JS逆向测试建图 Stage 4"
    props.author = "js-test-mapper"
    props.comments = "由 run-data 确定性投影生成"
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    build(args.input, args.output)
