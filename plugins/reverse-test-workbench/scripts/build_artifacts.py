#!/usr/bin/env python3
"""Build reverse-test-workbench DOCX/XLSX artifacts from canonical run data."""

from __future__ import annotations

import argparse
from datetime import datetime, timezone
import hashlib
import json
import os
from pathlib import Path
import shutil
import tempfile
import time
from typing import Any, Iterable

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    DOCX_AVAILABLE = True
    DOCX_IMPORT_ERROR = ""
except ImportError as exc:
    DOCX_AVAILABLE = False
    DOCX_IMPORT_ERROR = str(exc)

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    XLSX_AVAILABLE = True
    XLSX_IMPORT_ERROR = ""
except ImportError as exc:
    XLSX_AVAILABLE = False
    XLSX_IMPORT_ERROR = str(exc)

from validate_run_data import ValidationError, load_and_validate


DOCX_NAME = "过程小结.docx"
XLSX_NAME = "测试资产表.xlsx"
RUN_DATA_RELATIVE = Path("evidence/run-data.json")
STATE_RELATIVE = Path("evidence/_run-state.json")
BUILD_MANIFEST_RELATIVE = Path("evidence/_artifact-build.json")
STAGING_PREFIX = ".rtw-artifacts-"
STAGING_MARKER = ".rtw-staging.json"
STALE_STAGING_SECONDS = 86400
SHEET_SPECS: list[tuple[str, str, list[tuple[str, str]]]] = [
    (
        "01_批次与运行状态",
        "batches",
        [
            ("批次编号", "batch_id"), ("批次名称", "batch_name"),
            ("批次类型", "batch_type"), ("探索问题", "exploration_question"),
            ("运行定位", "run_positioning"), ("范围边界", "scope_boundary"),
            ("覆盖承诺", "coverage_commitment"), ("探索重点", "exploration_focus"),
            ("运行预算", "run_budget"), ("价值层级", "value_level"),
            ("信息增量摘要", "information_gain_summary"),
            ("已回答问题", "answered_questions"), ("未回答问题", "unanswered_questions"),
            ("结论用途", "conclusion_use"), ("执行器协议", "executor_protocol"),
            ("执行器版本", "executor_version"), ("浏览器通道", "browser_channel"),
            ("会话模式", "session_mode"), ("模型视觉能力", "model_vision_capability"),
            ("能力状态", "capability_state"), ("是否尝试修复", "repair_attempted"),
            ("执行器切换记录", "executor_switches"), ("计划来源", "plan_source"),
            ("是否计入当前计划", "in_current_plan"), ("状态", "status"),
            ("是否完成", "is_complete"), ("是否已处置", "is_disposed"),
            ("优先级说明", "priority_reason"), ("预计访问量", "estimated_visits"),
            ("实际到达页面/功能面", "actual_surfaces"),
            ("初步探索数", "initial_exploration_count"),
            ("有效探索数", "effective_exploration_count"),
            ("路径闭环数", "closed_path_count"), ("导航游标", "navigation_cursor"),
            ("返回游标", "return_cursor"), ("动态加入原因", "dynamic_reason"),
            ("未覆盖范围", "uncovered_scope"), ("探索债务", "exploration_debt"),
            ("开始时间", "started_at"), ("结束时间", "ended_at"),
            ("本批次耗时", "duration_ms"), ("阶段耗时摘要", "timing_summary"),
            ("阻塞原因", "blocker_reason"), ("下一步建议", "next_action"),
            ("关联小结章节", "report_section"), ("关联截图目录", "evidence_dir"),
        ],
    ),
    (
        "02_功能菜单清单",
        "navigation",
        [
            ("入口编号", "entry_id"), ("导航层级", "level"), ("一级导航", "level_1"),
            ("二级导航", "level_2"), ("三级导航", "level_3"),
            ("入口名称", "entry_name"), ("入口类型", "entry_type"),
            ("父入口", "parent_entry"), ("可见账号/角色", "visible_account"),
            ("入口地址/定位", "locator"), ("规范化页面或功能面", "normalized_surface"),
            ("入口分类", "entry_category"), ("盘点完整性", "inventory_completeness"),
            ("处置状态", "disposition_status"), ("本轮是否覆盖", "covered_this_run"),
            ("抽样依据", "sampling_basis"), ("未覆盖原因", "uncovered_reason"),
            ("待办/后续批次", "backlog_batch"), ("主要能力", "main_capability"),
            ("备注", "notes"),
        ],
    ),
    (
        "03_页面字段按钮清单",
        "page_inventory",
        [
            ("功能面编号", "surface_id"), ("页面/功能面名称", "surface_name"),
            ("所属入口", "entry_name"), ("核心对象", "core_object"), ("区域", "region"),
            ("字段/按钮/状态名称", "control_name"), ("类型", "control_type"),
            ("主要动作", "primary_action"), ("关键状态", "key_state"),
            ("可见条件", "visibility_condition"), ("行为指纹摘要", "behavior_fingerprint"),
            ("观察来源", "observation_source"), ("备注", "notes"),
        ],
    ),
    (
        "04_执行路径清单",
        "paths",
        [
            ("路径编号", "path_id"), ("批次编号", "batch_id"), ("路径名称", "path_name"),
            ("路径类型", "path_type"), ("流程状态", "flow_status"),
            ("探索问题", "exploration_question"), ("关系证据", "relationship_evidence"),
            ("前置条件", "preconditions"), ("实际步骤摘要", "steps_summary"),
            ("输入数据", "input_data"), ("状态前", "state_before"),
            ("状态后", "state_after"), ("观察结果", "observed_result"),
            ("断言强度", "assertion_strength"), ("结论类型", "conclusion_type"),
            ("阻塞点", "blocker"), ("返回游标", "return_cursor"),
            ("证据编号", "evidence_ids"),
        ],
    ),
    (
        "05_测试用例清单",
        "test_cases",
        [
            ("用例编号", "case_id"), ("批次编号", "batch_id"),
            ("模块/功能面", "surface"), ("用例标题", "title"),
            ("探索问题", "exploration_question"), ("前置条件", "preconditions"),
            ("测试步骤", "steps"), ("测试数据", "test_input"),
            ("预期/关系依据", "expected_basis"), ("断言强度", "assertion_strength"),
            ("预期来源", "expected_source"), ("实际结果", "actual_result"),
            ("结论", "conclusion"), ("证据编号", "evidence_ids"),
            ("是否适合自动化回归", "automation_candidate"),
        ],
    ),
    (
        "06_缺陷疑似问题清单",
        "issues",
        [
            ("问题编号", "issue_id"), ("批次编号", "batch_id"),
            ("模块/功能面", "surface"), ("标题", "title"), ("分类", "category"),
            ("严重级别", "severity"), ("首次线索", "first_signal"),
            ("复现/第二证据", "second_evidence"), ("实际结果", "actual_result"),
            ("判断依据", "judgment_basis"), ("预期来源", "expected_source"),
            ("问题来源", "issue_source"), ("影响范围", "impact_scope"),
            ("置信说明", "confidence"), ("证据编号", "evidence_ids"),
            ("状态", "status"), ("建议动作", "recommendation"), ("备注", "notes"),
        ],
    ),
    (
        "07_风险与待确认清单",
        "risks",
        [
            ("编号", "risk_id"), ("批次编号", "batch_id"),
            ("模块/功能面", "surface"), ("类型", "type"), ("优先级", "priority"),
            ("标题", "title"), ("描述", "description"), ("形成原因", "reason"),
            ("影响范围", "impact_scope"), ("关联路径/认知", "related_assets"),
            ("当前处理建议", "recommendation"), ("进入待办池原因", "backlog_reason"),
            ("状态", "status"),
        ],
    ),
    (
        "08_测试数据台账",
        "test_data",
        [
            ("数据编号", "data_id"), ("批次编号", "batch_id"),
            ("业务对象", "business_object"), ("数据标识", "data_identifier"),
            ("创建入口", "creation_entry"), ("创建账号/角色", "created_by"),
            ("创建时间", "created_at"), ("当前状态", "current_status"),
            ("关联路径/用例", "related_assets"), ("依赖关系", "dependencies"),
            ("是否建议清理", "cleanup_recommended"), ("清理建议", "cleanup_advice"),
            ("备注", "notes"),
        ],
    ),
    (
        "09_截图证据索引",
        "evidence",
        [
            ("证据编号", "evidence_id"), ("批次编号", "batch_id"),
            ("类型", "type"), ("关联对象", "related_object"), ("文件路径", "file_path"),
            ("图注", "caption"), ("截图时间", "captured_at"),
            ("页面/功能面位置", "location"), ("证明内容", "proves"),
            ("是否复用", "reused"), ("备注", "notes"),
        ],
    ),
    (
        "10_认知资产清单",
        "knowledge",
        [
            ("认知编号", "knowledge_id"), ("批次编号", "batch_id"),
            ("认知类型", "knowledge_type"), ("内容", "content"), ("来源", "source"),
            ("置信说明", "confidence"), ("是否可作为断言", "can_assert"),
            ("当前状态", "status"), ("替代/修订认知", "replacement"),
            ("受影响资产", "affected_assets"), ("关联页面/路径", "related_surface_path"),
            ("待确认问题", "confirmation_question"), ("影响/意义", "impact"),
            ("建议动作", "action"), ("结论用途", "conclusion_use"),
            ("更新记录", "change_log"),
        ],
    ),
]

NAVY = "24364B"
TEAL = "197278"
PALE = "E8F1F2"
LIGHT = "F4F7F8"
TEXT = "1F2933"
MUTED = "5E6C76"
WHITE = "FFFFFF"


def display(value: Any) -> Any:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, list):
        return "；".join(display(item) for item in value)
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False, separators=(",", ":"))
    return value


def _set_xlsx_value(cell: Any, value: Any) -> None:
    """Write report text as a literal value, never as an Excel formula."""
    shown = display(value)
    cell.value = shown
    if isinstance(shown, str) and shown.lstrip().startswith(("=", "+", "-", "@")):
        cell.data_type = "s"


def text(value: Any) -> str:
    shown = display(value)
    return str(shown) if shown != "" else ""


def _set_east_asia(run: Any, font_name: str) -> None:
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)


def _style_run(run: Any, size: float = 10.5, bold: bool = False, color: str = TEXT) -> None:
    run.font.name = "Arial"
    _set_east_asia(run, "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _add_label_paragraph(document: Document, label: str, value: Any) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    label_run = paragraph.add_run(f"{label}：")
    _style_run(label_run, bold=True, color=NAVY)
    value_run = paragraph.add_run(text(value) or "当前没有对应记录")
    _style_run(value_run)


def _add_list(document: Document, items: Iterable[Any], empty_text: str) -> None:
    materialized = list(items)
    if not materialized:
        paragraph = document.add_paragraph()
        _style_run(paragraph.add_run(empty_text), color=MUTED)
        return
    for item in materialized:
        paragraph = document.add_paragraph(style="List Bullet")
        if isinstance(item, dict):
            value = item.get("question") or item.get("title") or json.dumps(item, ensure_ascii=False)
            impact = item.get("impact")
            if impact:
                value = f"{value}（影响：{impact}）"
        else:
            value = item
        _style_run(paragraph.add_run(text(value)))


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for name, size, color in (
        ("Heading 1", 15, NAVY),
        ("Heading 2", 12, TEAL),
        ("Heading 3", 10.5, NAVY),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _add_title(document: Document, data: dict[str, Any]) -> None:
    run = data["run"]
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(4)
    title_run = paragraph.add_run(f"{run['system_name']} 探索测试过程小结")
    _style_run(title_run, size=20, bold=True, color=NAVY)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    _style_run(
        subtitle.add_run(
            f"运行编号 {run['run_id']}  |  状态 {run['status']}  |  目标 {run['target_url']}"
        ),
        size=9,
        color=MUTED,
    )


def _build_docx(data: dict[str, Any], path: Path, run_root: Path) -> None:
    document = Document()
    _configure_document(document)
    document.core_properties.title = f"{data['run']['system_name']} 探索测试过程小结"
    document.core_properties.subject = "reverse-test-workbench generated report"
    _add_title(document, data)

    run = data["run"]
    summary = data["summary"]
    document.add_heading("运行定位与一句话结论", level=1)
    _add_label_paragraph(document, "运行定位", run["run_positioning"])
    _add_label_paragraph(document, "范围边界", run["scope_boundary"])
    _add_label_paragraph(document, "覆盖承诺", run["coverage_commitment"])
    _add_label_paragraph(document, "探索重点", run["exploration_focus"])
    _add_label_paragraph(document, "运行预算", run["run_budget"])
    _add_label_paragraph(document, "结论", summary["one_line_conclusion"])

    document.add_heading("已回答问题", level=1)
    _add_list(document, summary["answered_questions"], "本轮尚未回答明确业务问题。")
    document.add_heading("未回答问题及影响", level=1)
    _add_list(document, summary["unanswered_questions"], "当前没有登记未回答问题。")

    document.add_heading("关键发现", level=1)
    findings = summary["key_findings"]
    if not findings:
        _add_list(document, [], "当前证据范围内没有形成关键发现。")
    for finding in findings:
        document.add_heading(text(finding.get("title", "未命名发现")), level=2)
        _add_label_paragraph(document, "观察事实", finding.get("fact"))
        _add_label_paragraph(document, "测试判断", finding.get("judgment"))
        _add_label_paragraph(document, "依据", finding.get("basis"))
        _add_label_paragraph(document, "影响", finding.get("impact"))
        _add_label_paragraph(document, "建议动作", finding.get("action"))
        _add_label_paragraph(document, "置信说明", finding.get("confidence"))

    document.add_heading("建议动作", level=1)
    _add_list(document, summary["recommendations"], "当前没有新增建议动作。")

    document.add_heading("批次摘要", level=1)
    batches = data["batches"]
    if not batches:
        _add_list(document, [], "当前没有批次记录。")
    else:
        table = document.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        headers = ("批次", "名称", "状态", "价值层级", "信息增量", "下一步")
        widths = (0.55, 1.15, 0.85, 0.8, 2.1, 1.35)
        for index, (header, width) in enumerate(zip(headers, widths)):
            cell = table.rows[0].cells[index]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _shade_cell(cell, NAVY)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _style_run(paragraph.add_run(header), size=8.5, bold=True, color=WHITE)
        for batch in batches:
            row = table.add_row().cells
            values = (
                batch.get("batch_id"), batch.get("batch_name"), batch.get("status"),
                batch.get("value_level"), batch.get("information_gain_summary"),
                batch.get("next_action"),
            )
            for index, (cell, value) in enumerate(zip(row, values)):
                cell.width = Inches(widths[index])
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                if len(table.rows) % 2 == 0:
                    _shade_cell(cell, LIGHT)
                paragraph = cell.paragraphs[0]
                _style_run(paragraph.add_run(text(value)), size=8.2)

    document.add_heading("覆盖审计", level=1)
    coverage = data["coverage"]
    coverage_items = [
        f"当前识别入口：{coverage.get('identified_entries', 0)}",
        f"实际到达功能面：{coverage.get('reached_surfaces', 0)}",
        f"初步探索：{coverage.get('initially_explored', 0)}；有效探索：{coverage.get('effectively_explored', 0)}；路径闭环：{coverage.get('closed_paths', 0)}",
        f"初始计划批次：{coverage.get('initial_batches', 0)}；动态新增：{coverage.get('dynamic_batches', 0)}；完成：{coverage.get('completed_batches', 0)}；处置：{coverage.get('disposed_batches', 0)}",
        f"待办池：{coverage.get('backlog_count', 0)}；未覆盖范围：{text(coverage.get('uncovered_scope', [])) or '无登记'}",
    ]
    _add_list(document, coverage_items, "当前没有覆盖审计数据。")

    included = [item for item in data["evidence"] if item.get("include_in_report")]
    if included:
        document.add_heading("代表证据", level=1)
        for item in included:
            relative = Path(str(item.get("file_path", "")))
            candidate = run_root / relative
            if candidate.is_file():
                try:
                    document.add_picture(str(candidate), width=Inches(6.3))
                    caption = document.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _style_run(caption.add_run(text(item.get("caption"))), size=8.5, color=MUTED)
                except Exception:
                    _add_label_paragraph(document, "证据未嵌入", item.get("file_path"))

    document.add_heading("结论边界", level=1)
    _add_label_paragraph(document, "适用范围", summary["conclusion_boundary"])
    footer = document.sections[-1].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_run(footer.add_run(f"reverse-test-workbench | {run['run_id']}"), size=8, color=MUTED)
    document.save(path)


def _enriched_batches(data: dict[str, Any]) -> list[dict[str, Any]]:
    run = data["run"]
    executor = run["executor"]
    common = {
        "run_positioning": run["run_positioning"],
        "scope_boundary": run["scope_boundary"],
        "coverage_commitment": run["coverage_commitment"],
        "exploration_focus": run["exploration_focus"],
        "run_budget": run["run_budget"],
        "executor_protocol": executor.get("protocol"),
        "executor_version": executor.get("version"),
        "browser_channel": executor.get("browser_channel"),
        "session_mode": executor.get("session_mode"),
        "model_vision_capability": executor.get("model_vision_capability"),
        "capability_state": executor.get("capability_state"),
        "repair_attempted": executor.get("repair_attempted", False),
        "executor_switches": executor.get("executor_switches", []),
    }
    return [{**common, **row} for row in data["batches"]]


def _sheet_rows(data: dict[str, Any], collection: str) -> list[dict[str, Any]]:
    if collection == "batches":
        return _enriched_batches(data)
    return data[collection]


def _width_for(header: str, values: list[Any]) -> float:
    lengths = [len(str(header)) * 2]
    for value in values[:200]:
        shown = text(value)
        lengths.append(max((len(line) for line in shown.splitlines()), default=0))
    width = max(lengths, default=10) + 2
    if any(token in header for token in ("编号", "状态", "类型", "层级", "是否", "时间")):
        return float(min(max(width, 12), 20))
    return float(min(max(width, 14), 36))


def _build_xlsx(data: dict[str, Any], path: Path) -> None:
    workbook = Workbook()
    workbook.remove(workbook.active)
    workbook.properties.title = f"{data['run']['system_name']} 测试资产表"
    workbook.properties.subject = "reverse-test-workbench generated assets"
    header_fill = PatternFill("solid", fgColor=NAVY)
    alternate_fill = PatternFill("solid", fgColor=LIGHT)
    header_font = Font(name="Arial", size=10, bold=True, color=WHITE)
    body_font = Font(name="Arial", size=9, color=TEXT)
    border = Border(bottom=Side(style="thin", color="D7DEE3"))

    for sheet_name, collection, columns in SHEET_SPECS:
        sheet = workbook.create_sheet(sheet_name)
        sheet.sheet_view.showGridLines = False
        sheet.freeze_panes = "A2"
        rows = _sheet_rows(data, collection)
        for column_index, (header, _) in enumerate(columns, start=1):
            cell = sheet.cell(1, column_index, header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border
        sheet.row_dimensions[1].height = 32

        for row_index, row in enumerate(rows, start=2):
            for column_index, (_, key) in enumerate(columns, start=1):
                cell = sheet.cell(row_index, column_index)
                _set_xlsx_value(cell, row.get(key, ""))
                cell.font = body_font
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = border
                if row_index % 2 == 0:
                    cell.fill = alternate_fill
            sheet.row_dimensions[row_index].height = 30

        last_row = max(sheet.max_row, 1)
        sheet.auto_filter.ref = f"A1:{get_column_letter(len(columns))}{last_row}"
        for column_index, (header, key) in enumerate(columns, start=1):
            values = [row.get(key, "") for row in rows]
            sheet.column_dimensions[get_column_letter(column_index)].width = _width_for(header, values)
        sheet.sheet_properties.pageSetUpPr.fitToPage = True
        sheet.page_setup.fitToWidth = 1
        sheet.page_setup.fitToHeight = 0
        sheet.auto_filter.ref = f"A1:{get_column_letter(len(columns))}{last_row}"

    workbook.save(path)


def _build_state_projection(data: dict[str, Any]) -> dict[str, Any]:
    run = data["run"]
    resume = run["resume"]
    executor = run["executor"]
    return {
        "schema_version": data["schema_version"],
        "run_id": run["run_id"],
        "system_name": run["system_name"],
        "target_url": run["target_url"],
        "status": run["status"],
        "run_positioning": run["run_positioning"],
        "scope_boundary": run["scope_boundary"],
        "coverage_commitment": run["coverage_commitment"],
        "exploration_focus": run["exploration_focus"],
        "run_budget": run["run_budget"],
        "budget_control": run["budget_control"],
        "account_context": run["account_context"],
        "executor_protocol": executor.get("protocol"),
        "executor_version": executor.get("version"),
        "browser_channel": executor.get("browser_channel"),
        "session_mode": executor.get("session_mode"),
        "model_vision_capability": executor.get("model_vision_capability"),
        "capability_state": executor.get("capability_state"),
        "artifact_preflight": run["artifact_preflight"],
        "timing": run["timing"],
        "current_batch": resume.get("current_batch"),
        "navigation_cursor": resume.get("navigation_cursor"),
        "return_cursor": resume.get("return_cursor"),
        "active_inserted_path": resume.get("active_inserted_path"),
        "backlog": resume.get("backlog", []),
        "last_known_good_page": resume.get("last_known_good_page", {}),
        "blockers": resume.get("blockers", []),
        "next_action": resume.get("next_action", ""),
        "coverage": data["coverage"],
    }


def _write_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def _validate_docx_structure(path: Path) -> dict[str, Any]:
    if not DOCX_AVAILABLE:
        raise RuntimeError(f"DOCX validation unavailable: {DOCX_IMPORT_ERROR}")
    document = Document(path)
    content = "\n".join(paragraph.text for paragraph in document.paragraphs)
    required = ("探索测试过程小结", "已回答问题", "未回答问题", "结论边界")
    missing = [fragment for fragment in required if fragment not in content]
    if missing:
        raise ValueError(f"DOCX structure missing sections: {', '.join(missing)}")
    if path.stat().st_size == 0:
        raise ValueError("DOCX output is empty")
    return {
        "status": "passed",
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "required_sections": list(required),
    }


def _validate_xlsx_structure(path: Path) -> dict[str, Any]:
    if not XLSX_AVAILABLE:
        raise RuntimeError(f"XLSX validation unavailable: {XLSX_IMPORT_ERROR}")
    workbook = load_workbook(path, read_only=True, data_only=False)
    expected_sheets = [name for name, _, _ in SHEET_SPECS]
    if workbook.sheetnames != expected_sheets:
        workbook.close()
        raise ValueError(f"XLSX sheet mismatch: {workbook.sheetnames}")
    checked_headers: dict[str, int] = {}
    for sheet_name, _, columns in SHEET_SPECS:
        sheet = workbook[sheet_name]
        actual = [sheet.cell(1, index).value for index in range(1, len(columns) + 1)]
        expected = [header for header, _ in columns]
        if actual != expected:
            workbook.close()
            raise ValueError(f"XLSX header mismatch in {sheet_name}")
        checked_headers[sheet_name] = len(expected)
    workbook.close()
    if path.stat().st_size == 0:
        raise ValueError("XLSX output is empty")
    return {
        "status": "passed",
        "sheets": expected_sheets,
        "header_counts": checked_headers,
    }


def _cleanup_stale_staging(
    output_dir: Path,
    max_age_seconds: int = STALE_STAGING_SECONDS,
) -> list[str]:
    output_root = output_dir.resolve()
    removed: list[str] = []
    if not output_root.exists():
        return removed
    now = time.time()
    for candidate in output_root.iterdir():
        if not candidate.is_dir() or not candidate.name.startswith(STAGING_PREFIX):
            continue
        resolved = candidate.resolve()
        if resolved.parent != output_root:
            continue
        marker = resolved / STAGING_MARKER
        if not marker.is_file():
            continue
        try:
            marker_data = json.loads(marker.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            continue
        if marker_data.get("owner") != "reverse-test-workbench":
            continue
        if now - marker.stat().st_mtime < max_age_seconds:
            continue
        try:
            shutil.rmtree(resolved)
        except OSError:
            continue
        else:
            removed.append(candidate.name)
    return removed


def _publish_transaction(
    staged_targets: dict[Path, Path],
    backup_dir: Path,
) -> None:
    backup_dir.mkdir(parents=True, exist_ok=True)
    backups: dict[Path, Path] = {}
    published: list[Path] = []
    for index, target in enumerate(staged_targets.values()):
        if target.exists():
            if not target.is_file():
                raise OSError(f"artifact target is not a file: {target}")
            backup = backup_dir / f"{index:02d}-{target.name}"
            shutil.copy2(target, backup)
            backups[target] = backup
    try:
        for source, target in staged_targets.items():
            target.parent.mkdir(parents=True, exist_ok=True)
            os.replace(source, target)
            published.append(target)
    except Exception:
        for target in published:
            if target.exists() and target.is_file():
                target.unlink()
        for target, backup in backups.items():
            if backup.exists():
                target.parent.mkdir(parents=True, exist_ok=True)
                os.replace(backup, target)
        raise


def _artifact_decision(
    data: dict[str, Any],
    artifact: str,
    explicitly_skipped: bool,
    import_available: bool,
    import_error: str,
) -> tuple[bool, str]:
    if explicitly_skipped:
        return False, "explicitly_skipped"
    if data["run"]["artifact_preflight"].get(f"{artifact}_generation") == "unavailable":
        return False, "preflight_unavailable"
    if not import_available:
        return False, f"runtime_import_unavailable: {import_error}"
    return True, ""


def _data_fingerprint(data: dict[str, Any]) -> str:
    encoded = json.dumps(
        data,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def _file_fingerprint(path: Path) -> dict[str, Any]:
    content = path.read_bytes()
    return {"size_bytes": len(content), "sha256": hashlib.sha256(content).hexdigest()}


def _matches_file_fingerprint(path: Path, expected: Any) -> bool:
    if not path.is_file() or not isinstance(expected, dict):
        return False
    try:
        actual = _file_fingerprint(path)
    except OSError:
        return False
    return actual == {
        "size_bytes": expected.get("size_bytes"),
        "sha256": expected.get("sha256"),
    }


def _load_existing_manifest(output_dir: Path) -> dict[str, Any] | None:
    path = output_dir / BUILD_MANIFEST_RELATIVE
    if not path.is_file():
        return None
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    return value if isinstance(value, dict) else None


def _unchanged_result(
    data: dict[str, Any],
    output_dir: Path,
    request: dict[str, bool],
    input_sha256: str,
) -> dict[str, Any] | None:
    manifest = _load_existing_manifest(output_dir)
    if not manifest:
        return None
    if manifest.get("input_sha256") != input_sha256 or manifest.get("requested") != request:
        return None
    artifacts = manifest.get("artifacts", {})
    file_integrity = manifest.get("file_integrity", {})
    required_files = {
        RUN_DATA_RELATIVE.as_posix(): output_dir / RUN_DATA_RELATIVE,
        STATE_RELATIVE.as_posix(): output_dir / STATE_RELATIVE,
    }
    for name, requested in request.items():
        if not requested:
            continue
        artifact = artifacts.get(name, {})
        if artifact.get("status") == "generated":
            target = output_dir / (DOCX_NAME if name == "docx" else XLSX_NAME)
            required_files[target.relative_to(output_dir).as_posix()] = target
            if not _matches_file_fingerprint(target, artifact.get("integrity")):
                return None
    for relative, target in required_files.items():
        if not _matches_file_fingerprint(target, file_integrity.get(relative)):
            return None
    result: dict[str, Any] = {
        "status": "unchanged",
        "previous_status": manifest.get("status"),
        "manifest": str(output_dir / BUILD_MANIFEST_RELATIVE),
        "run_data": str(output_dir / RUN_DATA_RELATIVE),
        "run_state": str(output_dir / STATE_RELATIVE),
        "generated": manifest.get("generated", []),
        "skipped": manifest.get("skipped", []),
    }
    if artifacts.get("docx", {}).get("status") == "generated":
        result["docx"] = str(output_dir / DOCX_NAME)
    if artifacts.get("xlsx", {}).get("status") == "generated":
        result["xlsx"] = str(output_dir / XLSX_NAME)
    return result


def build(
    data: dict[str, Any],
    output_dir: Path,
    state_only: bool = False,
    skip_docx: bool = False,
    skip_xlsx: bool = False,
) -> dict[str, Any]:
    started = time.perf_counter()
    output_dir.mkdir(parents=True, exist_ok=True)
    evidence_dir = output_dir / "evidence"
    evidence_dir.mkdir(parents=True, exist_ok=True)
    run_root = output_dir
    stale_removed = _cleanup_stale_staging(output_dir)
    request = {
        "docx": not state_only and not skip_docx,
        "xlsx": not state_only and not skip_xlsx,
    }
    input_sha256 = _data_fingerprint(data)
    if not state_only:
        unchanged = _unchanged_result(data, output_dir, request, input_sha256)
        if unchanged is not None:
            unchanged["stale_staging_removed"] = stale_removed
            return unchanged
    temp_dir = Path(tempfile.mkdtemp(prefix=STAGING_PREFIX, dir=output_dir))
    _write_json(
        temp_dir / STAGING_MARKER,
        {
            "owner": "reverse-test-workbench",
            "run_id": data["run"]["run_id"],
            "created_at": datetime.now(timezone.utc).isoformat(),
        },
    )
    try:
        temp_docx = temp_dir / DOCX_NAME
        temp_xlsx = temp_dir / XLSX_NAME
        temp_run_data = temp_dir / RUN_DATA_RELATIVE
        temp_state = temp_dir / STATE_RELATIVE
        temp_manifest = temp_dir / BUILD_MANIFEST_RELATIVE
        _write_json(temp_run_data, data)
        _write_json(temp_state, _build_state_projection(data))

        staged_targets: dict[Path, Path] = {
            temp_run_data: output_dir / RUN_DATA_RELATIVE,
            temp_state: output_dir / STATE_RELATIVE,
        }
        artifacts = {
            "docx": {
                "status": "not_requested" if state_only else "pending",
                "path": DOCX_NAME,
                "reason": "state_only" if state_only else "",
                "preserved_previous": False,
                "structure_validation": "not_run",
            },
            "xlsx": {
                "status": "not_requested" if state_only else "pending",
                "path": XLSX_NAME,
                "reason": "state_only" if state_only else "",
                "preserved_previous": False,
                "structure_validation": "not_run",
            },
        }
        validations: dict[str, Any] = {}

        if not state_only:
            docx_enabled, docx_reason = _artifact_decision(
                data, "docx", skip_docx, DOCX_AVAILABLE, DOCX_IMPORT_ERROR
            )
            xlsx_enabled, xlsx_reason = _artifact_decision(
                data, "xlsx", skip_xlsx, XLSX_AVAILABLE, XLSX_IMPORT_ERROR
            )
            if docx_enabled:
                _build_docx(data, temp_docx, run_root)
                validations["docx"] = _validate_docx_structure(temp_docx)
                artifacts["docx"].update(
                    status="generated", structure_validation="passed"
                )
                staged_targets[temp_docx] = output_dir / DOCX_NAME
            else:
                artifacts["docx"].update(
                    status="skipped",
                    reason=docx_reason,
                    preserved_previous=(output_dir / DOCX_NAME).is_file(),
                )
            if xlsx_enabled:
                _build_xlsx(data, temp_xlsx)
                validations["xlsx"] = _validate_xlsx_structure(temp_xlsx)
                artifacts["xlsx"].update(
                    status="generated", structure_validation="passed"
                )
                staged_targets[temp_xlsx] = output_dir / XLSX_NAME
            else:
                artifacts["xlsx"].update(
                    status="skipped",
                    reason=xlsx_reason,
                    preserved_previous=(output_dir / XLSX_NAME).is_file(),
                )

        file_integrity = {
            RUN_DATA_RELATIVE.as_posix(): _file_fingerprint(temp_run_data),
            STATE_RELATIVE.as_posix(): _file_fingerprint(temp_state),
        }
        for artifact_name, temp_path in (("docx", temp_docx), ("xlsx", temp_xlsx)):
            if artifacts[artifact_name]["status"] == "generated":
                integrity = _file_fingerprint(temp_path)
                artifacts[artifact_name]["integrity"] = integrity
                file_integrity[artifacts[artifact_name]["path"]] = integrity

        generated_heavy = [
            name for name in ("docx", "xlsx") if artifacts[name]["status"] == "generated"
        ]
        skipped_heavy = [
            name for name in ("docx", "xlsx") if artifacts[name]["status"] == "skipped"
        ]
        if state_only:
            status = "state_updated"
            mode = "state_only"
        elif skipped_heavy:
            status = "partial"
            mode = "full"
        else:
            status = "generated"
            mode = "full"

        preflight = data["run"]["artifact_preflight"]
        visual_status = (
            "skipped_structure_fallback"
            if preflight.get("visual_render_unavailable")
            else "pending_external_render"
        )
        manifest = {
            "schema_version": "1.0",
            "run_id": data["run"]["run_id"],
            "built_at": datetime.now(timezone.utc).isoformat(),
            "mode": mode,
            "status": status,
            "input_sha256": input_sha256,
            "requested": request,
            "artifacts": artifacts,
            "file_integrity": file_integrity,
            "generated": generated_heavy,
            "skipped": skipped_heavy,
            "validations": validations,
            "visual_validation": {
                "status": visual_status,
                "libreoffice": preflight.get("libreoffice", "unknown"),
                "fallback": preflight.get("fallback", "none"),
                "claimed_passed": False,
            },
            "cleanup": {
                "stale_staging_removed": stale_removed,
                "active_staging_cleanup": "automatic",
            },
            "duration_ms": int((time.perf_counter() - started) * 1000),
        }
        _write_json(temp_manifest, manifest)
        staged_targets[temp_manifest] = output_dir / BUILD_MANIFEST_RELATIVE
        _publish_transaction(staged_targets, temp_dir / "backup")

        outputs: dict[str, Any] = {
            "status": status,
            "manifest": str(output_dir / BUILD_MANIFEST_RELATIVE),
            "run_data": str(output_dir / RUN_DATA_RELATIVE),
            "run_state": str(output_dir / STATE_RELATIVE),
            "generated": generated_heavy,
            "skipped": skipped_heavy,
        }
        if artifacts["docx"]["status"] == "generated":
            outputs["docx"] = str(output_dir / DOCX_NAME)
        if artifacts["xlsx"]["status"] == "generated":
            outputs["xlsx"] = str(output_dir / XLSX_NAME)
        return outputs
    finally:
        if temp_dir.exists():
            shutil.rmtree(temp_dir, ignore_errors=True)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path, help="Canonical run-data JSON")
    parser.add_argument("--output-dir", required=True, type=Path, help="Run output directory")
    parser.add_argument(
        "--state-only",
        action="store_true",
        help="Only update evidence/run-data.json and evidence/_run-state.json",
    )
    parser.add_argument("--skip-docx", action="store_true", help="Do not update DOCX")
    parser.add_argument("--skip-xlsx", action="store_true", help="Do not update XLSX")
    args = parser.parse_args()
    try:
        data = load_and_validate(args.input)
        outputs = build(
            data,
            args.output_dir,
            state_only=args.state_only,
            skip_docx=args.skip_docx,
            skip_xlsx=args.skip_xlsx,
        )
    except (ValidationError, OSError, RuntimeError, ValueError) as exc:
        print(f"artifact generation failed: {exc}")
        return 1
    print(json.dumps(outputs, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
